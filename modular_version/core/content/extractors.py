from __future__ import annotations
from pathlib import Path

from .limits import MAX_DOCX_SIZE_MB, MAX_XLSX_SIZE_MB, MAX_PDF_SIZE_MB, MAX_EXTRACTED_TEXT_CHARS
from .models import ExtractResult
from .pdf_content import extract_pdf_document

def _limit_bytes(mb: int) -> int:
    return mb * 1024 * 1024

def extract_docx(path: str, *, stop_flag=None) -> ExtractResult:
    if Path(path).stat().st_size > _limit_bytes(MAX_DOCX_SIZE_MB):
        return ExtractResult(text='', status='TOO_LARGE', source='docx', error_text='DOCX exceeds size limit')
    try:
        from docx import Document
    except Exception as e:
        return ExtractResult(text='', status='ERROR', source='docx', error_text=f'python-docx import failed: {e}')
    doc = Document(path)
    parts = []
    for p in doc.paragraphs:
        if stop_flag and stop_flag.get('stop'):
            raise InterruptedError('Stopped by user')
        txt = (p.text or '').strip()
        if txt:
            parts.append(txt)
    for table in doc.tables:
        if stop_flag and stop_flag.get('stop'):
            raise InterruptedError('Stopped by user')
        for row in table.rows:
            row_text = ' | '.join((cell.text or '').strip() for cell in row.cells if (cell.text or '').strip())
            if row_text:
                parts.append(row_text)
    text = '\n'.join(parts).strip()[:MAX_EXTRACTED_TEXT_CHARS]
    return ExtractResult(text=text, status='OK' if text else 'NO_TEXT', source='docx', quality_score=1.0 if text else 0.0, text_length=len(text))

def extract_xlsx(path: str, *, stop_flag=None) -> ExtractResult:
    if Path(path).stat().st_size > _limit_bytes(MAX_XLSX_SIZE_MB):
        return ExtractResult(text='', status='TOO_LARGE', source='xlsx', error_text='XLSX exceeds size limit')
    try:
        from openpyxl import load_workbook
    except Exception as e:
        return ExtractResult(text='', status='ERROR', source='xlsx', error_text=f'openpyxl import failed: {e}')
    wb = load_workbook(path, read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        if stop_flag and stop_flag.get('stop'):
            raise InterruptedError('Stopped by user')
        parts.append(f'[SHEET] {ws.title}')
        for row in ws.iter_rows(values_only=True):
            if stop_flag and stop_flag.get('stop'):
                raise InterruptedError('Stopped by user')
            vals = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if vals:
                parts.append(' | '.join(vals))
    text = '\n'.join(parts).strip()[:MAX_EXTRACTED_TEXT_CHARS]
    return ExtractResult(text=text, status='OK' if text else 'NO_TEXT', source='xlsx', quality_score=1.0 if text else 0.0, text_length=len(text))

def extract_content(path: str, ext: str, *, stop_flag=None) -> ExtractResult:
    ext = (ext or '').lower().lstrip('.')
    if ext == 'docx':
        return extract_docx(path, stop_flag=stop_flag)
    if ext == 'xlsx':
        return extract_xlsx(path, stop_flag=stop_flag)
    if ext == 'pdf':
        p = Path(path)
        if p.stat().st_size > _limit_bytes(MAX_PDF_SIZE_MB):
            return ExtractResult(text='', status='TOO_LARGE', source='pdf', error_text='PDF exceeds size limit')
        return extract_pdf_document(path, stop_flag=stop_flag)
    return ExtractResult(text='', status='UNSUPPORTED', source=ext or 'unknown', error_text='Unsupported extension')
